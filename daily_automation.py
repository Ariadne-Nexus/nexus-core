#!/usr/bin/env python3
"""Daily automation script for summarizing notes and posting updates."""

from __future__ import annotations

import argparse
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Optional

import requests
from docx import Document
from openai import OpenAI
from openpyxl import Workbook, load_workbook

BASE_DIR = Path(__file__).resolve().parent
NOTES_DIR = BASE_DIR / "notes"
OUTPUT_DIR = BASE_DIR / "output"
TXT_PATH = OUTPUT_DIR / "daily_summary.txt"
DOCX_PATH = OUTPUT_DIR / "daily_summary.docx"
XLSX_PATH = OUTPUT_DIR / "daily_log.xlsx"


# Logging --------------------------------------------------------------------
def log(message: str) -> None:
    """Simple timestamped logger."""
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print(f"[{now}] {message}", flush=True)


# CLI ------------------------------------------------------------------------
def parse_args(argv: Optional[list[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Daily automation runner (v2)")
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Run without writing files or touching GitHub; log intended actions only.",
    )
    return parser.parse_args(argv)


# OpenAI ---------------------------------------------------------------------
def check_openai_client() -> OpenAI:
    """Validate the OpenAI environment and return a ready client."""
    log("Checking OPENAI_API_KEY...")
    if not os.getenv("OPENAI_API_KEY"):
        log("ERROR: OPENAI_API_KEY is not set.")
        raise SystemExit(1)

    try:
        client = OpenAI()
    except Exception as exc:  # pragma: no cover - depends on environment
        log(f"ERROR: Failed to initialize OpenAI client: {exc}")
        raise SystemExit(1)

    log("OpenAI client initialized.")
    return client


def summarize_notes(client: OpenAI, notes_text: str) -> str:
    """Request a concise summary of the provided notes."""
    log("Step 1: Summarizing notes via OpenAI...")
    prompt = f"""
You are my daily automation assistant.

Summarize the following notes into:
- A 2–3 sentence high-level summary
- 3–5 bullet points of key actions or insights

Respond in plain text, no markdown headings. Notes:
{notes_text}
""".strip()

    try:
        response = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt,
        )
        piece = response.output[0].content[0].text
        summary_text = getattr(piece, "value", str(piece))
        log("Got structured summary from OpenAI.")
        return summary_text.strip()
    except Exception as exc:  # pragma: no cover - depends on environment
        log(f"ERROR during summarization: {exc}")
        raise


# Notes ----------------------------------------------------------------------
def load_recent_notes(max_files: int = 5) -> str:
    """Load the most recent text/markdown notes for summarization."""
    log("Loading recent notes...")
    if not NOTES_DIR.exists():
        log(f"No notes directory found at {NOTES_DIR}, using placeholder text.")
        return "No notes were found; this is a placeholder run."

    note_files = sorted(
        [p for p in NOTES_DIR.iterdir() if p.suffix in {".txt", ".md"}],
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )

    if not note_files:
        log("No .txt or .md files found in notes; using placeholder text.")
        return "No notes were found; this is a placeholder run."

    selected = note_files[:max_files]
    log(f"Using {len(selected)} note file(s):")
    for file_path in selected:
        log(f"  - {file_path.name}")

    chunks: list[str] = []
    for path in selected:
        try:
            chunks.append(f"# {path.name}\n")
            chunks.append(path.read_text(encoding="utf-8"))
            chunks.append("\n\n")
        except Exception as exc:  # pragma: no cover - depends on environment
            log(f"WARNING: Could not read {path}: {exc}")

    combined = "\n".join(chunks).strip()
    if not combined:
        return "Notes existed but could not be read; this is a placeholder run."
    return combined


# Output helpers -------------------------------------------------------------
def ensure_output_dir() -> None:
    OUTPUT_DIR.mkdir(exist_ok=True)


def write_txt_summary(summary: str) -> Path:
    log("Writing text summary...")
    ensure_output_dir()
    TXT_PATH.write_text(summary + "\n", encoding="utf-8")
    log(f"Text summary written to {TXT_PATH}")
    return TXT_PATH


def write_docx_summary(summary: str) -> Path:
    log("Writing DOCX summary...")
    ensure_output_dir()

    doc = Document()
    today = datetime.now().strftime("%Y-%m-%d")
    doc.add_heading(f"Daily Summary – {today}", level=1)
    doc.add_paragraph(summary)
    doc.save(DOCX_PATH)

    log(f"DOCX summary written to {DOCX_PATH}")
    return DOCX_PATH


def append_excel_log(summary: str) -> Path:
    log("Appending to Excel log...")
    ensure_output_dir()

    if XLSX_PATH.exists():
        wb = load_workbook(XLSX_PATH)
        ws = wb.active
    else:
        wb = Workbook()
        ws = wb.active
        ws.title = "Daily Log"
        ws.append(["Timestamp", "Summary"])

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    ws.append([timestamp, summary])
    wb.save(XLSX_PATH)

    log(f"Excel log updated at {XLSX_PATH}")
    return XLSX_PATH


# GitHub integration ---------------------------------------------------------
def create_or_update_github_task(summary: str) -> Optional[str]:
    """Create or update a GitHub issue with the daily summary.

    Returns the URL of the issue if successful, otherwise ``None``.
    """

    repo = os.getenv("GITHUB_REPO")
    token = os.getenv("GITHUB_TOKEN")
    if not repo or not token:
        log("GITHUB_REPO or GITHUB_TOKEN not set; skipping GitHub task update.")
        return None

    headers = {
        "Accept": "application/vnd.github+json",
        "Authorization": f"Bearer {token}",
        "X-GitHub-Api-Version": "2022-11-28",
    }

    today = datetime.now().strftime("%Y-%m-%d")
    title = f"Daily Automation – {today}"
    body = f"Daily automation summary for {today}.\n\n{summary}"

    try:
        issues_response = requests.get(
            f"https://api.github.com/repos/{repo}/issues",
            params={"state": "open", "per_page": 100},
            headers=headers,
            timeout=10,
        )
        issues_response.raise_for_status()
        issues = issues_response.json()

        for issue in issues:
            if issue.get("title") == title:
                issue_number = issue.get("number")
                update_response = requests.patch(
                    f"https://api.github.com/repos/{repo}/issues/{issue_number}",
                    headers=headers,
                    json={"body": body},
                    timeout=10,
                )
                update_response.raise_for_status()
                log(f"Updated existing GitHub issue #{issue_number}.")
                return issue.get("html_url")

        create_response = requests.post(
            f"https://api.github.com/repos/{repo}/issues",
            headers=headers,
            json={"title": title, "body": body},
            timeout=10,
        )
        create_response.raise_for_status()
        issue = create_response.json()
        log(f"Created new GitHub issue #{issue.get('number')}.")
        return issue.get("html_url")

    except requests.RequestException as exc:  # pragma: no cover - network behavior
        log(f"ERROR while updating GitHub: {exc}")
        return None


# Runner ---------------------------------------------------------------------
def run(dry_run: bool) -> int:
    log(f"=== Daily v2 automation run starting (dry_run={dry_run}) ===")

    client = check_openai_client()
    notes_text = load_recent_notes()
    summary = summarize_notes(client, notes_text)

    if dry_run:
        log("DRY-RUN: Skipping file writes and GitHub issue.")
        log("DRY-RUN: Would write outputs to:")
        log(f"  TXT:  {TXT_PATH}")
        log(f"  DOCX: {DOCX_PATH}")
        log(f"  XLSX: {XLSX_PATH}")
        today = datetime.now().strftime("%Y-%m-%d")
        log("DRY-RUN: Would create/update a GitHub issue titled:")
        log(f"  'Daily Automation – {today}' in repo: {os.getenv('GITHUB_REPO')}")
        log("=== Daily v2 automation dry-run finished successfully ===")
        return 0

    txt_path = write_txt_summary(summary)
    docx_path = write_docx_summary(summary)
    xlsx_path = append_excel_log(summary)
    issue_url = create_or_update_github_task(summary)

    log("Run complete.")
    log(
        "Outputs:\n"
        f"  TXT:  {txt_path}\n"
        f"  DOCX: {docx_path}\n"
        f"  XLSX: {xlsx_path}"
    )
    if issue_url:
        log(f"GitHub task: {issue_url}")
    else:
        log("GitHub task: skipped or failed (see logs above).")

    log("=== Daily v2 automation run finished successfully ===")
    return 0


def main(argv: Optional[list[str]] = None) -> int:
    args = parse_args(argv)
    return run(dry_run=args.dry_run)


if __name__ == "__main__":
    sys.exit(main())
